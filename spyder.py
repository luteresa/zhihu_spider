#coding='utf-8'
from bs4 import *
from urllib3 import *
import re

disable_warnings()

http = PoolManager()
html_zhihu='https://www.zhihu.com/people/matongxue/answers'
html_zhihu_matongxue='https://www.zhihu.com/question/23149768/answer/282842210'
def getContent(url):
    r = http.request('GET',url)
    soup = BeautifulSoup(r.data,'lxml')
    filename=soup.title.string[:-5]
    #print(soup.prettify())
    def filter(tag):
        if tag.has_attr('class') and True:#tag.class=='List-item':
            return True
        return False

    import docx
    from docx.shared import Inches
    file = docx.Document()

    for tag  in soup.find_all(True):
    #for tag  in soup.find_all(filter):
    # for tag in soup.find_all('<div',class_=='List-item')
        if tag.name == 'p' or tag.name=='li':
            if tag.parent.name == 'span':
                print(tag.text)

            file.add_paragraph(tag.text)
        if tag.name=='img':
            print(tag.get('class'))
            if tag.get('class') == ['Avatar', 'AuthorInfo-avatar'] or \
                    tag.get('class') == ['Avatar','Avatar--large UserLink-avatar']:
                pass
            else:
                link=tag.get('src')

                if link.endswith(('.jpg', '.png')):
                    print(link)
                    print('width:{},height:{}'.format(tag.get('data-rawwidth'), tag.get('data-rawheight')))
                    print('==========')
                    r2 = http.request('GET', link)
                    f = open('t3.jpg', 'wb')
                    f.write(r2.data)
                    f.close()
                    file.add_picture('t3.jpg',width=Inches(5))

    file.save('{}.docx'.format(filename))

getContent(html_zhihu_matongxue)